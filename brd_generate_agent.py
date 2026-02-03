from langgraph.graph import StateGraph, START, END
from typing import TypedDict
from google import genai
from docx import Document
from htmldocx import HtmlToDocx
import os
import re
<<<<<<< HEAD
from docx import Document
from fpdf import FPDF
=======
from datetime import datetime
>>>>>>> 95f6ff8 (Updated code with new version)
from dotenv import load_dotenv
import textwrap

load_dotenv()

<<<<<<< HEAD
# ---------------- STATE DEFINITION ----------------
=======

>>>>>>> 95f6ff8 (Updated code with new version)
class BRDState(TypedDict):
    project_name: str
    user_input: str
<<<<<<< HEAD
    brd_template_file: str
    is_valid: bool
    final_pdf: str

# ---------------- NODE 1 — INPUT ----------------
def input_node(state: BRDState) -> BRDState:
    print(f"📌 Project Input Received: {state['user_input']}")
    return state

# ---------------- HELPER: PARSE TEMPLATE ----------------
def parse_template_headings(template_file: str):
    doc = Document(template_file)
    headings = []
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            headings.append(para.text.strip())
    return headings

# ---------------- NODE 2 — GENERATE BRD ----------------
def validate_and_generate_node(state: BRDState) -> BRDState:
    GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
    if not GOOGLE_API_KEY:
        raise ValueError("GOOGLE_API_KEY not found")

    client = genai.Client(api_key=GOOGLE_API_KEY)
    model = "gemini-2.5-flash"

    project_name = state["project_name"]
    project_description = state["user_input"]
    template_file = state["brd_template_file"]

    if not template_file:
        raise ValueError("No BRD template file provided.")

    headings = parse_template_headings(template_file)

    # Combine all headings in one prompt
    prompt_headings = "\n".join([f"{i+1}. {h}" for i, h in enumerate(headings)])

    prompt = f"""
You are a professional Business Analyst.

Generate a complete BRD for the following project:

PROJECT NAME: {project_name}
PROJECT DESCRIPTION: {project_description}

Please generate detailed sections for all of these headings:
{prompt_headings}

Rules:
- Keep it professional.
- Keep sections concise but informative.
- Output plain text only, start each section with its heading as shown.
"""

    # Single API call
    response = client.models.generate_content(model=model, contents=prompt)
    brd_text = response.text.strip()

    # Optional: split into sections
    sections = {}
    current_heading = None
    for line in brd_text.split("\n"):
        line = line.strip()
        if any(line.startswith(f"{i+1}. {h}") for i, h in enumerate(headings)):
            current_heading = line
            sections[current_heading] = ""
        elif current_heading:
            sections[current_heading] += line + "\n"

    state["brd_sections"] = sections
    state["brd_text"] = brd_text
    state["is_valid"] = True
    print("✅ BRD Generated Successfully based on template!")
    return state

# ---------------- NODE 3 — PARSE SECTIONS ----------------
def split_sections_node(state: BRDState) -> BRDState:
    brd_lines = state["brd_text"].split("\n")
    sections = {}
    current_heading = None

    for line in brd_lines:
        line = line.strip()
        if re.match(r"^\d+(\.\d+)*\s", line):
            current_heading = line
            sections[current_heading] = ""
        elif current_heading:
            sections[current_heading] += line + "\n"

    state["brd_sections"] = sections
    print("📂 BRD split into sections successfully!")
    return state

# ---------------- PDF CREATION ----------------
def markdown_to_pdf(brd_text: str, output_pdf: str = "final_brd.pdf") -> str:
    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)

    # Set proper margins
    left_margin = 15
    right_margin = 15
    pdf.set_left_margin(left_margin)
    pdf.set_right_margin(right_margin)

    pdf.add_page()

    page_width = pdf.w - left_margin - right_margin  # Safe writable width

    # Title
    pdf.set_font("Arial", "B", 16)
    pdf.cell(page_width, 10, "Business Requirements Document (BRD)", ln=True, align="C")
    pdf.ln(8)

    lines = brd_text.split("\n")

    for line in lines:
        line = line.strip()

        if not line:
            pdf.ln(3)
            continue

        # Section Headings (1. Introduction, 2. Scope, etc.)
        if re.match(r"^\d+(\.\d+)*\s", line):
            pdf.set_font("Arial", "B", 13)
            pdf.ln(2)
            pdf.multi_cell(page_width, 8, line)
            pdf.ln(1)

        # Normal Paragraph Text
        else:
            pdf.set_font("Arial", "", 11)

            # Prevent long unbroken text from breaking layout
            safe_line = line.encode("latin-1", "replace").decode("latin-1")

            pdf.multi_cell(page_width, 6, safe_line)
            pdf.ln(1)

    pdf.output(output_pdf)
    print(f"📕 Final PDF Ready: {output_pdf}")
    return output_pdf

# ---------------- NODE 4 — OUTPUT PDF ----------------
def output_node(state: BRDState) -> BRDState:
    pdf_path = "final_brd.pdf"
    markdown_to_pdf(state["brd_text"], pdf_path)
    state["final_pdf"] = pdf_path   # ⭐ THIS is what UI needs
    return state


# ---------------- GRAPH ----------------
builder = StateGraph(BRDState)
builder.add_node("input_node", input_node)
builder.add_node("validate_and_generate_node", validate_and_generate_node)
builder.add_node("split_sections_node", split_sections_node)
builder.add_node("output_node", output_node)

builder.add_edge(START, "input_node")
builder.add_edge("input_node", "validate_and_generate_node")
=======
    headings: list
    brd_html: str
    output_path: str
    file_name: str
    final_docx: str
    is_valid: bool
    brd_template_file: str


def extract_headings_node(state: BRDState) -> BRDState:
    """Extract Heading 1 from the template"""
    doc = Document(state["brd_template_file"])
    headings = []
    for para in doc.paragraphs:
        if para.style.name == "Heading 1":
            clean = re.sub(r"\s*\(.*?\)\s*", "", para.text.strip())
            if clean:
                headings.append(clean)
    state["headings"] = headings
    print("Extracted Headings:", headings)
    return state

def validate_input_node(state: BRDState) -> BRDState:
    """Check if user input is valid"""
    state["is_valid"] = len(state["user_input"].split()) >= 5
    if not state["is_valid"]:
        print("Invalid input detected")
    return state

def generate_brd_html_node(state: BRDState) -> BRDState:
    """Generate HTML BRD using Google Gemini API"""
    GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
    if not GOOGLE_API_KEY:
        raise Exception("GOOGLE_API_KEY not set")

    client = genai.Client(api_key=GOOGLE_API_KEY)
    model = "gemini-3-flash-preview"

    prompt = f"""
You are an expert Business Analyst. Generate a professional BRD for the following project:

Project Description: {state['user_input']}

Use the following main headings: {', '.join(state['headings'])}.
Include numbered subheadings, bullet points, and tables where necessary.
Output HTML using <h1>, <h2>, <h3>, <p>, <ul>, <li>, <table>, <tr>, <th>, <td>.
Do not use markdown.
Tables must include headers.
"""
    response = client.models.generate_content(model=model, contents=prompt)
    state["brd_html"] = getattr(response, "text", None) or \
                        getattr(response, "output_text", None) or \
                        response.contents[0].text
    print("AI HTML Generated")
    return state

def html_to_word_node(state: BRDState) -> BRDState:
    """Convert HTML BRD to Word document"""
    if not state.get("brd_html"):
        print("ERROR: brd_html is empty!")
        return state

    os.makedirs("files", exist_ok=True)

    doc = Document()
    parser = HtmlToDocx()
    parser.add_html_to_document(state["brd_html"], doc)

    for table in doc.tables:
        table.style = "Table Grid"

    safe_proj_name = re.sub(r'[\\/*?:"<>|]', "_", state["project_name"])
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    file_name = f"BRD_{safe_proj_name}_{timestamp}.docx"
    output_path = os.path.join("files", file_name)

    doc.save(output_path)

    state["output_path"] = os.path.abspath(output_path)
    state["file_name"] = file_name
    state["final_docx"] = state["output_path"]
    print(f"Word document : {state['final_docx']}")
    return state

def invalid_output_node(state: BRDState) -> BRDState:
    print("Please provide a valid project description.")
    return state


builder = StateGraph(BRDState)
builder.add_node("extract_headings_node", extract_headings_node)
builder.add_node("validate_input_node", validate_input_node)
builder.add_node("generate_brd_html_node", generate_brd_html_node)
builder.add_node("html_to_word_node", html_to_word_node)
builder.add_node("invalid_output_node", invalid_output_node)

builder.add_edge(START, "extract_headings_node")
builder.add_edge("extract_headings_node", "validate_input_node")
>>>>>>> 95f6ff8 (Updated code with new version)

def route_after_validation(state: BRDState):
    return "generate_brd_html_node" if state["is_valid"] else "invalid_output_node"

builder.add_conditional_edges("validate_input_node", route_after_validation)
builder.add_edge("generate_brd_html_node", "html_to_word_node")
builder.add_edge("html_to_word_node", END)
builder.add_edge("invalid_output_node", END)

graph = builder.compile()
